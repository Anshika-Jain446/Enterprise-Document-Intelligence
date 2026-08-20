import io
import os
import json
from pathlib import Path

import pandas as pd
import streamlit as st
from werkzeug.security import generate_password_hash, check_password_hash

from chunking import ChunkingEngine
from model import GeminiModel

try:
    import psycopg2
    from psycopg2 import sql
    from psycopg2.extras import RealDictCursor, Json
except ImportError:
    psycopg2 = None
    sql = None
    RealDictCursor = None
    Json = None

try:
    from pgvector.psycopg2 import register_vector
except ImportError:
    register_vector = None


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="Enterprise Document Intelligence",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# CONFIG
# ============================================================

DATABASE_URL = os.getenv("DATABASE_URL", os.getenv("POSTGRESQL_URL", ""))
PG_HOST = os.getenv("PG_HOST", "localhost")
PG_PORT = os.getenv("PG_PORT", "5432")
PG_DATABASE = os.getenv("PG_DATABASE", "enterprise_rag")
PG_USER = os.getenv("PG_USER", "postgres")
PG_PASSWORD = os.getenv("PG_PASSWORD", "postgres")

# Recommended: set ADMIN_USERNAME to your username in Streamlit secrets/env.
# If it is not set and the database has no admin yet, the first existing user
# is promoted to admin so an existing installation does not become inaccessible.
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "").strip()

# pgvector / embedding configuration.
PGVECTOR_ENABLED = os.getenv("PGVECTOR_ENABLED", "true").lower() in {"1", "true", "yes", "on"}
PGVECTOR_DIMENSION = int(os.getenv("PGVECTOR_DIMENSION", "768"))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "gemini-embedding-001")


# ============================================================
# POSTGRESQL STORE
# ============================================================

class PostgreSQLStore:
    def __init__(self):
        self.connection = None

    def connect(self):
        if psycopg2 is None:
            raise RuntimeError(
                "psycopg2 is not installed. Add psycopg2-binary to requirements.txt."
            )

        if self.connection is not None:
            try:
                if self.connection.closed == 0:
                    try:
                        self.connection.rollback()
                    except Exception:
                        pass
                    return self.connection
            except Exception:
                pass
            try:
                self.connection.close()
            except Exception:
                pass
            self.connection = None

        if DATABASE_URL:
            self.connection = psycopg2.connect(DATABASE_URL, connect_timeout=15)
        else:
            self.connection = psycopg2.connect(
                host=PG_HOST,
                port=PG_PORT,
                database=PG_DATABASE,
                user=PG_USER,
                password=PG_PASSWORD,
                connect_timeout=15,
            )

        self.connection.autocommit = False
        if register_vector is not None:
            try:
                register_vector(self.connection)
            except Exception:
                pass
        cursor = self.connection.cursor()
        try:
            cursor.execute("SET search_path TO public")
            self.connection.commit()
        except Exception:
            self.connection.rollback()
            raise
        finally:
            cursor.close()
        return self.connection

    def rollback(self):
        if self.connection is not None:
            try:
                self.connection.rollback()
            except Exception:
                pass

    def close(self):
        if self.connection is not None:
            try:
                self.connection.close()
            except Exception:
                pass
            self.connection = None

    def column_exists(self, cursor, table_name, column_name):
        cursor.execute(
            """
            SELECT EXISTS (
                SELECT 1 FROM information_schema.columns
                WHERE table_schema = 'public'
                  AND table_name = %s
                  AND column_name = %s
            )
            """,
            (table_name, column_name),
        )
        return bool(cursor.fetchone()[0])

    def add_column_if_missing(self, cursor, table_name, column_name, definition):
        cursor.execute(
            sql.SQL("ALTER TABLE public.{} ADD COLUMN IF NOT EXISTS {} {}").format(
                sql.Identifier(table_name),
                sql.Identifier(column_name),
                sql.SQL(definition),
            )
        )

    def initialize(self):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            cursor.execute("SET search_path TO public")

            # --------------------------------------------------------
            # USERS + ROLES
            # --------------------------------------------------------
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS public.users (
                    id SERIAL PRIMARY KEY,
                    username VARCHAR(255),
                    password_hash TEXT,
                    role VARCHAR(50) DEFAULT 'user',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            self.add_column_if_missing(cursor, "users", "username", "VARCHAR(255)")
            self.add_column_if_missing(cursor, "users", "password_hash", "TEXT")
            self.add_column_if_missing(cursor, "users", "role", "VARCHAR(50) DEFAULT 'user'")
            self.add_column_if_missing(cursor, "users", "created_at", "TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
            conn.commit()

            # Normalize missing/invalid roles.
            cursor.execute(
                """
                UPDATE public.users
                SET role = 'user'
                WHERE role IS NULL OR LOWER(role) NOT IN ('admin', 'user')
                """
            )
            conn.commit()

            # Keep usernames unique where possible.
            try:
                cursor.execute(
                    """
                    CREATE UNIQUE INDEX IF NOT EXISTS users_username_unique_idx
                    ON public.users(username)
                    """
                )
                conn.commit()
            except Exception:
                conn.rollback()

            # --------------------------------------------------------
            # DOCUMENTS
            # --------------------------------------------------------
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS public.documents (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER,
                    filename TEXT,
                    file_type TEXT,
                    file_size BIGINT,
                    file_data BYTEA,
                    chunking_method TEXT,
                    chunk_size INTEGER,
                    chunk_overlap INTEGER,
                    metadata JSONB DEFAULT '{}'::jsonb,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            for name, definition in {
                "user_id": "INTEGER",
                "filename": "TEXT",
                "file_type": "TEXT",
                "file_size": "BIGINT",
                "file_data": "BYTEA",
                "chunking_method": "TEXT",
                "chunk_size": "INTEGER",
                "chunk_overlap": "INTEGER",
                "metadata": "JSONB DEFAULT '{}'::jsonb",
                "created_at": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            }.items():
                self.add_column_if_missing(cursor, "documents", name, definition)
            conn.commit()

            # --------------------------------------------------------
            # CHUNKS
            # --------------------------------------------------------
            # The embedding column is added when pgvector is available.
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS public.document_chunks (
                    id SERIAL PRIMARY KEY,
                    document_id INTEGER,
                    chunk_id INTEGER,
                    chunk_type TEXT,
                    content TEXT,
                    page INTEGER,
                    tokens INTEGER,
                    characters INTEGER,
                    metadata JSONB DEFAULT '{}'::jsonb,
                    chunk_data JSONB DEFAULT '{}'::jsonb,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            conn.commit()

            for name, definition in {
                "document_id": "INTEGER",
                "chunk_id": "INTEGER",
                "chunk_type": "TEXT",
                "content": "TEXT",
                "page": "INTEGER",
                "tokens": "INTEGER",
                "characters": "INTEGER",
                "metadata": "JSONB DEFAULT '{}'::jsonb",
                "chunk_data": "JSONB DEFAULT '{}'::jsonb",
                "created_at": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            }.items():
                self.add_column_if_missing(cursor, "document_chunks", name, definition)
            conn.commit()

            if PGVECTOR_ENABLED:
                try:
                    cursor.execute("CREATE EXTENSION IF NOT EXISTS vector")
                    conn.commit()
                    self.add_column_if_missing(
                        cursor, "document_chunks", "embedding",
                        f"vector({PGVECTOR_DIMENSION})"
                    )
                    conn.commit()
                    cursor.execute(
                        """
                        CREATE INDEX IF NOT EXISTS idx_document_chunks_embedding_hnsw
                        ON public.document_chunks
                        USING hnsw (embedding vector_cosine_ops)
                        """
                    )
                    conn.commit()
                except Exception:
                    conn.rollback()

            # --------------------------------------------------------
            # CONVERSATIONS / MESSAGES
            # --------------------------------------------------------
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS public.conversations (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER,
                    title TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS public.messages (
                    id SERIAL PRIMARY KEY,
                    conversation_id INTEGER,
                    role TEXT,
                    content TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            for name, definition in {
                "user_id": "INTEGER",
                "title": "TEXT",
                "created_at": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            }.items():
                self.add_column_if_missing(cursor, "conversations", name, definition)
            for name, definition in {
                "conversation_id": "INTEGER",
                "role": "TEXT",
                "content": "TEXT",
                "created_at": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
            }.items():
                self.add_column_if_missing(cursor, "messages", name, definition)
            conn.commit()

            # --------------------------------------------------------
            # REMOVE OLD FKs before orphan cleanup, then recreate.
            # --------------------------------------------------------
            for table, constraint in [
                ("documents", "documents_user_id_fkey"),
                ("document_chunks", "document_chunks_document_id_fkey"),
                ("conversations", "conversations_user_id_fkey"),
                ("messages", "messages_conversation_id_fkey"),
            ]:
                try:
                    cursor.execute(
                        sql.SQL("ALTER TABLE public.{} DROP CONSTRAINT IF EXISTS {}").format(
                            sql.Identifier(table), sql.Identifier(constraint)
                        )
                    )
                    conn.commit()
                except Exception:
                    conn.rollback()

            # Remove bad legacy rows before adding FKs.
            cleanup = [
                """
                DELETE FROM public.documents d
                WHERE d.user_id IS NOT NULL
                  AND NOT EXISTS (SELECT 1 FROM public.users u WHERE u.id = d.user_id)
                """,
                """
                DELETE FROM public.document_chunks c
                WHERE c.document_id IS NOT NULL
                  AND NOT EXISTS (SELECT 1 FROM public.documents d WHERE d.id = c.document_id)
                """,
                """
                DELETE FROM public.conversations c
                WHERE c.user_id IS NOT NULL
                  AND NOT EXISTS (SELECT 1 FROM public.users u WHERE u.id = c.user_id)
                """,
                """
                DELETE FROM public.messages m
                WHERE m.conversation_id IS NOT NULL
                  AND NOT EXISTS (
                      SELECT 1 FROM public.conversations c WHERE c.id = m.conversation_id
                  )
                """,
            ]
            for statement in cleanup:
                try:
                    cursor.execute(statement)
                    conn.commit()
                except Exception:
                    conn.rollback()

            # Add relationships if possible.
            for statement in [
                """
                ALTER TABLE public.documents
                ADD CONSTRAINT documents_user_id_fkey
                FOREIGN KEY (user_id) REFERENCES public.users(id) ON DELETE CASCADE
                """,
                """
                ALTER TABLE public.document_chunks
                ADD CONSTRAINT document_chunks_document_id_fkey
                FOREIGN KEY (document_id) REFERENCES public.documents(id) ON DELETE CASCADE
                """,
                """
                ALTER TABLE public.conversations
                ADD CONSTRAINT conversations_user_id_fkey
                FOREIGN KEY (user_id) REFERENCES public.users(id) ON DELETE CASCADE
                """,
                """
                ALTER TABLE public.messages
                ADD CONSTRAINT messages_conversation_id_fkey
                FOREIGN KEY (conversation_id) REFERENCES public.conversations(id) ON DELETE CASCADE
                """,
            ]:
                try:
                    cursor.execute(statement)
                    conn.commit()
                except Exception:
                    conn.rollback()

            # --------------------------------------------------------
            # INDEXES
            # --------------------------------------------------------
            index_statements = [
                "CREATE INDEX IF NOT EXISTS idx_users_role ON public.users(role)",
                "CREATE UNIQUE INDEX IF NOT EXISTS documents_user_filename_unique_idx ON public.documents(user_id, filename)",
                "CREATE UNIQUE INDEX IF NOT EXISTS document_chunks_document_chunk_unique_idx ON public.document_chunks(document_id, chunk_id)",
                "CREATE INDEX IF NOT EXISTS idx_document_chunks_document ON public.document_chunks(document_id)",
                """
                CREATE INDEX IF NOT EXISTS idx_document_chunks_content
                ON public.document_chunks USING gin (
                    to_tsvector('english', COALESCE(content, ''))
                )
                """,
            ]
            for statement in index_statements:
                try:
                    cursor.execute(statement)
                    conn.commit()
                except Exception:
                    conn.rollback()

            # --------------------------------------------------------
            # ADMIN MIGRATION
            # --------------------------------------------------------
            # Explicit ADMIN_USERNAME always wins.
            if ADMIN_USERNAME:
                cursor.execute(
                    "UPDATE public.users SET role = 'admin' WHERE LOWER(username) = LOWER(%s)",
                    (ADMIN_USERNAME,),
                )
                conn.commit()

            cursor.execute("SELECT COUNT(*) FROM public.users WHERE role = 'admin'")
            admin_count = cursor.fetchone()[0]

            if admin_count == 0:
                cursor.execute("SELECT id FROM public.users ORDER BY id LIMIT 1")
                first_user = cursor.fetchone()
                if first_user:
                    cursor.execute(
                        "UPDATE public.users SET role = 'admin' WHERE id = %s",
                        (first_user[0],),
                    )
                    conn.commit()

            # Final required-column verification.
            required = {
                "users": ["id", "username", "password_hash", "role", "created_at"],
                "documents": [
                    "id", "user_id", "filename", "file_type", "file_size", "file_data",
                    "chunking_method", "chunk_size", "chunk_overlap", "metadata", "created_at",
                ],
                "document_chunks": [
                    "id", "document_id", "chunk_id", "chunk_type", "content", "page",
                    "tokens", "characters", "metadata", "chunk_data", "created_at",
                ],
                "conversations": ["id", "user_id", "title", "created_at"],
                "messages": ["id", "conversation_id", "role", "content", "created_at"],
            }
            missing = []
            for table, columns in required.items():
                for column in columns:
                    if not self.column_exists(cursor, table, column):
                        missing.append(f"{table}.{column}")
            if missing:
                raise RuntimeError("PostgreSQL migration failed. Missing columns: " + ", ".join(missing))

            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    # ========================================================
    # USER / ADMIN METHODS
    # ========================================================

    def create_user(self, username, password, role="user"):
        username = username.strip()
        if not username or not password:
            return None
        role = role if role in {"admin", "user"} else "user"
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            cursor.execute(
                """
                INSERT INTO public.users (username, password_hash, role)
                VALUES (%s, %s, %s)
                RETURNING id, username, role
                """,
                (username, generate_password_hash(password), role),
            )
            result = cursor.fetchone()
            conn.commit()
            return result
        except Exception:
            conn.rollback()
            return None
        finally:
            if cursor:
                cursor.close()

    def authenticate(self, username, password):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            cursor.execute(
                """
                SELECT id, username, password_hash, role, created_at
                FROM public.users
                WHERE LOWER(username) = LOWER(%s)
                """,
                (username.strip(),),
            )
            user = cursor.fetchone()
            conn.rollback()
            if not user:
                return None
            if not check_password_hash(user["password_hash"], password):
                return None
            return dict(user)
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    def get_users(self):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            cursor.execute(
                """
                SELECT
                    u.id,
                    u.username,
                    u.role,
                    u.created_at,
                    COUNT(DISTINCT d.id) AS document_count,
                    COUNT(DISTINCT c.id) AS conversation_count
                FROM public.users u
                LEFT JOIN public.documents d ON d.user_id = u.id
                LEFT JOIN public.conversations c ON c.user_id = u.id
                GROUP BY u.id
                ORDER BY u.id
                """
            )
            rows = [dict(r) for r in cursor.fetchall()]
            conn.rollback()
            return rows
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    def set_user_role(self, user_id, role):
        if role not in {"admin", "user"}:
            return False
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            if role == "user":
                cursor.execute("SELECT COUNT(*) FROM public.users WHERE role = 'admin'")
                if cursor.fetchone()[0] <= 1:
                    conn.rollback()
                    return False
            cursor.execute("UPDATE public.users SET role = %s WHERE id = %s", (role, user_id))
            changed = cursor.rowcount > 0
            conn.commit()
            return changed
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    def delete_user(self, user_id):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT role FROM public.users WHERE id = %s", (user_id,))
            row = cursor.fetchone()
            if not row:
                conn.rollback()
                return False, "User not found."
            if row[0] == "admin":
                cursor.execute("SELECT COUNT(*) FROM public.users WHERE role = 'admin'")
                if cursor.fetchone()[0] <= 1:
                    conn.rollback()
                    return False, "You cannot delete the last admin."
            cursor.execute("DELETE FROM public.users WHERE id = %s", (user_id,))
            conn.commit()
            return cursor.rowcount > 0, ""
        except Exception as exc:
            conn.rollback()
            return False, str(exc)
        finally:
            if cursor:
                cursor.close()

    def get_admin_stats(self):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            queries = {
                "users": "SELECT COUNT(*) FROM public.users",
                "admins": "SELECT COUNT(*) FROM public.users WHERE role = 'admin'",
                "documents": "SELECT COUNT(*) FROM public.documents",
                "chunks": "SELECT COUNT(*) FROM public.document_chunks",
                "conversations": "SELECT COUNT(*) FROM public.conversations",
                "messages": "SELECT COUNT(*) FROM public.messages",
            }
            result = {}
            for key, query in queries.items():
                cursor.execute(query)
                result[key] = cursor.fetchone()[0]
            conn.rollback()
            return result
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    # ========================================================
    # DOCUMENTS / CHUNKS
    # ========================================================

    def save_document(self, user_id, filename, file_type, file_bytes,
                      chunking_method, chunk_size, chunk_overlap, metadata):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            cursor.execute(
                "DELETE FROM public.documents WHERE user_id = %s AND filename = %s",
                (user_id, filename),
            )
            cursor.execute(
                """
                INSERT INTO public.documents (
                    user_id, filename, file_type, file_size, file_data,
                    chunking_method, chunk_size, chunk_overlap, metadata
                )
                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                RETURNING id
                """,
                (
                    user_id, filename, file_type, len(file_bytes), psycopg2.Binary(file_bytes),
                    chunking_method, chunk_size, chunk_overlap, Json(metadata or {}),
                ),
            )
            document_id = cursor.fetchone()[0]
            conn.commit()
            return document_id
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    def save_chunks(self, document_id, chunks, embedder=None):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            for index, chunk in enumerate(chunks):
                if not isinstance(chunk, dict):
                    chunk = {"content": str(chunk)}
                chunk_id = chunk.get("chunk_id", index)
                metadata = chunk.get("metadata", {})
                content = str(chunk.get("content", ""))
                embedding = None

                if PGVECTOR_ENABLED and embedder and content.strip():
                    try:
                        embedding = embedder(content)
                        if embedding is not None:
                            if len(embedding) != PGVECTOR_DIMENSION:
                                raise ValueError(
                                    f"Embedding dimension {len(embedding)} != PGVECTOR_DIMENSION {PGVECTOR_DIMENSION}"
                                )
                    except Exception:
                        embedding = None

                if PGVECTOR_ENABLED:
                    cursor.execute(
                        """
                        INSERT INTO public.document_chunks (
                            document_id, chunk_id, chunk_type, content, page,
                            tokens, characters, metadata, chunk_data, embedding
                        )
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                        ON CONFLICT (document_id, chunk_id)
                        DO UPDATE SET
                            chunk_type = EXCLUDED.chunk_type,
                            content = EXCLUDED.content,
                            page = EXCLUDED.page,
                            tokens = EXCLUDED.tokens,
                            characters = EXCLUDED.characters,
                            metadata = EXCLUDED.metadata,
                            chunk_data = EXCLUDED.chunk_data,
                            embedding = EXCLUDED.embedding
                        """,
                        (document_id, chunk_id, chunk.get("chunk_type"), content,
                         chunk.get("page"), chunk.get("tokens"), chunk.get("characters"),
                         Json(metadata), Json(chunk),
                         json.dumps(embedding) if embedding is not None else None),
                    )
                else:
                    cursor.execute(
                        """
                        INSERT INTO public.document_chunks (
                            document_id, chunk_id, chunk_type, content, page,
                            tokens, characters, metadata, chunk_data
                        )
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                        ON CONFLICT (document_id, chunk_id)
                        DO UPDATE SET
                            chunk_type = EXCLUDED.chunk_type,
                            content = EXCLUDED.content,
                            page = EXCLUDED.page,
                            tokens = EXCLUDED.tokens,
                            characters = EXCLUDED.characters,
                            metadata = EXCLUDED.metadata,
                            chunk_data = EXCLUDED.chunk_data
                        """,
                        (document_id, chunk_id, chunk.get("chunk_type"), content,
                         chunk.get("page"), chunk.get("tokens"), chunk.get("characters"),
                         Json(metadata), Json(chunk)),
                    )
            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    def get_documents(self, user_id, is_admin=False):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            if is_admin:
                cursor.execute(
                    """
                    SELECT d.id, d.user_id, u.username, d.filename, d.file_type,
                           d.file_size, d.chunking_method, d.chunk_size, d.chunk_overlap,
                           d.metadata, d.created_at
                    FROM public.documents d
                    LEFT JOIN public.users u ON u.id = d.user_id
                    ORDER BY d.created_at DESC
                    """
                )
            else:
                cursor.execute(
                    """
                    SELECT id, user_id, filename, file_type, file_size,
                           chunking_method, chunk_size, chunk_overlap, metadata, created_at
                    FROM public.documents
                    WHERE user_id = %s
                    ORDER BY created_at DESC
                    """,
                    (user_id,),
                )
            rows = [dict(r) for r in cursor.fetchall()]
            conn.rollback()
            return rows
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    # ========================================================
    # RETRIEVAL
    # ========================================================

    def search_chunks(self, user_id, query, selected_document_ids=None,
                      chunk_types=None, top_k=5, retrieval_method="PostgreSQL Full-Text + LIKE",
                      query_embedding=None):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            conditions = ["d.user_id = %s"]
            base_params = [user_id]

            if selected_document_ids:
                conditions.append("d.id = ANY(%s)")
                base_params.append(selected_document_ids)
            if chunk_types:
                conditions.append("c.chunk_type = ANY(%s)")
                base_params.append(chunk_types)

            vector_ready = PGVECTOR_ENABLED and query_embedding is not None

            if retrieval_method in {"pgvector Semantic", "Hybrid (pgvector + PostgreSQL Full-Text)"} and vector_ready:
                vector_literal = json.dumps(query_embedding)
                if retrieval_method == "pgvector Semantic":
                    sql_query = f"""
                        SELECT c.id, c.document_id, c.chunk_id, c.chunk_type, c.content,
                               c.page, c.tokens, c.characters, c.metadata, c.chunk_data,
                               d.filename, d.chunking_method,
                               1 - (c.embedding <=> %s::vector) AS similarity_score
                        FROM public.document_chunks c
                        JOIN public.documents d ON d.id = c.document_id
                        WHERE {' AND '.join(conditions)}
                          AND c.embedding IS NOT NULL
                        ORDER BY c.embedding <=> %s::vector
                        LIMIT %s
                    """
                    cursor.execute(sql_query, [vector_literal] + base_params + [vector_literal, top_k])
                else:
                    sql_query = f"""
                        SELECT c.id, c.document_id, c.chunk_id, c.chunk_type, c.content,
                               c.page, c.tokens, c.characters, c.metadata, c.chunk_data,
                               d.filename, d.chunking_method,
                               (0.7 * (1 - (c.embedding <=> %s::vector)) +
                                0.3 * ts_rank(
                                    to_tsvector('english', COALESCE(c.content, '')),
                                    plainto_tsquery('english', %s)
                                )) AS similarity_score
                        FROM public.document_chunks c
                        JOIN public.documents d ON d.id = c.document_id
                        WHERE {' AND '.join(conditions)}
                          AND c.embedding IS NOT NULL
                        ORDER BY similarity_score DESC
                        LIMIT %s
                    """
                    cursor.execute(sql_query, [vector_literal, query] + base_params + [top_k])
            else:
                fts = "to_tsvector('english', COALESCE(c.content, '')) @@ plainto_tsquery('english', %s)"
                like = "LOWER(COALESCE(c.content, '')) LIKE LOWER(%s)"
                if retrieval_method == "PostgreSQL Full-Text":
                    conditions.append(fts)
                    match_params = [query]
                elif retrieval_method == "LIKE Keyword":
                    conditions.append(like)
                    match_params = [f"%{query}%"]
                else:
                    conditions.append(f"({fts} OR {like})")
                    match_params = [query, f"%{query}%"]

                sql_query = f"""
                    SELECT c.id, c.document_id, c.chunk_id, c.chunk_type, c.content,
                           c.page, c.tokens, c.characters, c.metadata, c.chunk_data,
                           d.filename, d.chunking_method,
                           ts_rank(to_tsvector('english', COALESCE(c.content, '')),
                                   plainto_tsquery('english', %s)) AS similarity_score
                    FROM public.document_chunks c
                    JOIN public.documents d ON d.id = c.document_id
                    WHERE {' AND '.join(conditions)}
                    ORDER BY similarity_score DESC, c.id DESC
                    LIMIT %s
                """
                cursor.execute(sql_query, [query] + base_params + match_params + [top_k])

            rows = [dict(r) for r in cursor.fetchall()]
            conn.rollback()
            for row in rows:
                row["source"] = row.get("filename")
                row["search_type"] = retrieval_method
                row["source_type"] = "table" if row.get("chunk_type") == "Table" else "document"
            return rows
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    # ========================================================
    # CONVERSATIONS
    # ========================================================

    def create_conversation(self, user_id, title="New Conversation"):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO public.conversations (user_id, title) VALUES (%s,%s) RETURNING id",
                (user_id, title),
            )
            cid = cursor.fetchone()[0]
            conn.commit()
            return cid
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()

    def save_message(self, conversation_id, role, content):
        conn = self.connect()
        cursor = None
        try:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO public.messages (conversation_id, role, content) VALUES (%s,%s,%s)",
                (conversation_id, role, content),
            )
            conn.commit()
        except Exception:
            conn.rollback()
            raise
        finally:
            if cursor:
                cursor.close()


# ============================================================
# FILE EXTRACTION
# ============================================================

def extract_file(uploaded_file):
    filename = uploaded_file.name
    suffix = Path(filename).suffix.lower()
    file_bytes = uploaded_file.getvalue()
    text = ""
    pages, tables, images, visuals = [], [], [], []
    metadata = {
        "uploaded_file_name": filename,
        "filename": filename,
        "file_type": suffix,
        "file_size": len(file_bytes),
    }

    if suffix in {".txt", ".md", ".csv"}:
        text = file_bytes.decode("utf-8", errors="ignore")
        if suffix == ".csv":
            try:
                dataframe = pd.read_csv(io.BytesIO(file_bytes))
                tables.append({
                    "table": dataframe.fillna("").to_dict(orient="records"),
                    "page": 1,
                    "source": filename,
                })
            except Exception:
                pass

    elif suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            page_texts = []
            for page_number, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
                pages.append({"page": page_number, "text": page_text})
            text = "\n\n".join(page_texts)
            metadata["page_count"] = len(reader.pages)
        except Exception as exc:
            st.warning(f"PDF extraction failed: {exc}")

    elif suffix == ".docx":
        try:
            from docx import Document
            document = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            for table in document.tables:
                tables.append({
                    "table": [[cell.text for cell in row.cells] for row in table.rows],
                    "page": 1,
                    "source": filename,
                })
        except Exception as exc:
            st.warning(f"DOCX extraction failed: {exc}")

    elif suffix == ".xlsx":
        try:
            workbook = pd.ExcelFile(io.BytesIO(file_bytes))
            for sheet_name in workbook.sheet_names:
                dataframe = pd.read_excel(workbook, sheet_name=sheet_name)
                tables.append({
                    "table": dataframe.fillna("").to_dict(orient="records"),
                    "sheet": sheet_name,
                    "page": sheet_name,
                    "source": filename,
                })
                text += f"\n\nSheet: {sheet_name}\n{dataframe.to_string(index=False)}"
        except Exception as exc:
            st.warning(f"XLSX extraction failed: {exc}")

    return {
        "text": text,
        "metadata": metadata,
        "tables": tables,
        "images": images,
        "visuals": visuals,
        "pages": pages,
        "source": filename,
        "filename": filename,
        "file_name": filename,
    }


# ============================================================
# CHUNKING — INGESTION, NOT RETRIEVAL
# ============================================================

CHUNKING_METHODS = {
    "Recursive": "recursive_chunking",
    "Character": "character_chunking",
    "Token": "token_chunking",
    "Markdown": "markdown_chunking",
    "Contextual": "contextual_chunking",
    "Table": "table_chunking",
    "Image": "image_chunking",
    "Visual": "visual_chunking",
    "Multimodal": "multimodal_chunking",
}


def create_chunks(extracted, method, chunk_size, chunk_overlap):
    engine = ChunkingEngine(
        text=extracted.get("text", ""),
        metadata=extracted.get("metadata", {}),
        tables=extracted.get("tables", []),
        images=extracted.get("images", []),
        visuals=extracted.get("visuals", []),
        pages=extracted.get("pages", []),
        source=extracted.get("source", "Unknown"),
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    method_name = CHUNKING_METHODS[method]
    chunk_function = getattr(engine, method_name)
    return chunk_function(), engine


# ============================================================
# SESSION STATE
# ============================================================

def initialize_state():
    defaults = {
        "authenticated": False,
        "user": None,
        "db": None,
        "llm": None,
        "conversation_id": None,
        "messages": [],
        "documents": [],
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


initialize_state()


# ============================================================
# DATABASE INITIALIZATION
# ============================================================

if st.session_state.db is None:
    try:
        st.session_state.db = PostgreSQLStore()
        st.session_state.db.initialize()
    except Exception as exc:
        try:
            st.session_state.db.rollback()
        except Exception:
            pass
        st.error("PostgreSQL database initialization failed.")
        st.exception(exc)
        st.stop()


# ============================================================
# LOGIN / REGISTER
# ============================================================

if not st.session_state.authenticated:
    st.title("📄 Enterprise Document Intelligence")
    st.caption("PostgreSQL-backed Agentic RAG")

    login_tab, register_tab = st.tabs(["🔐 Login", "📝 Register"])

    with login_tab:
        username = st.text_input("Username", key="login_username")
        password = st.text_input("Password", type="password", key="login_password")

        if st.button("Login", use_container_width=True):
            if not username or not password:
                st.warning("Enter username and password.")
            else:
                try:
                    user = st.session_state.db.authenticate(username, password)
                    if user:
                        st.session_state.authenticated = True
                        st.session_state.user = user
                        st.session_state.conversation_id = st.session_state.db.create_conversation(
                            user["id"], "Enterprise RAG"
                        )
                        st.session_state.messages = []
                        st.rerun()
                    else:
                        st.error("Invalid username or password.")
                except Exception as exc:
                    st.session_state.db.rollback()
                    st.error(f"Login failed: {exc}")

    with register_tab:
        new_username = st.text_input("Username", key="register_username")
        new_password = st.text_input("Password", type="password", key="register_password")
        confirm_password = st.text_input("Confirm password", type="password", key="register_confirm")

        if st.button("Create account", use_container_width=True):
            if not new_username or not new_password:
                st.warning("Username and password are required.")
            elif new_password != confirm_password:
                st.error("Passwords do not match.")
            else:
                user = st.session_state.db.create_user(new_username, new_password, "user")
                if user:
                    st.success("Account created. You can now log in.")
                else:
                    st.error("Username already exists or registration failed.")

    st.stop()


# ============================================================
# MODEL
# ============================================================

if st.session_state.llm is None:
    try:
        st.session_state.llm = GeminiModel()
    except Exception as exc:
        st.error(f"LLM initialization failed: {exc}")
        st.stop()


# ============================================================
# ROLE HELPERS
# ============================================================

current_user = st.session_state.user
is_admin = str(current_user.get("role", "user")).lower() == "admin"


# ============================================================
# HEADER
# ============================================================

st.title("📄 Enterprise Document Intelligence")
st.caption(
    "PostgreSQL • Agentic RAG • Separate Chunking & Retrieval • Multimodal Documents"
)


# ============================================================
# ADMIN PANEL
# ============================================================

if is_admin:
    with st.expander("🛡️ Admin Panel", expanded=False):
        st.subheader("System administration")
        st.caption("Admin controls are separate from the normal user document controls.")

        try:
            stats = st.session_state.db.get_admin_stats()
            a, b, c, d, e, f = st.columns(6)
            a.metric("Users", stats["users"])
            b.metric("Admins", stats["admins"])
            c.metric("Documents", stats["documents"])
            d.metric("Chunks", stats["chunks"])
            e.metric("Conversations", stats["conversations"])
            f.metric("Messages", stats["messages"])
        except Exception as exc:
            st.warning(f"Could not load admin statistics: {exc}")

        st.divider()
        st.subheader("👥 Users")

        try:
            users = st.session_state.db.get_users()
        except Exception as exc:
            users = []
            st.error(f"Could not load users: {exc}")

        if users:
            for user in users:
                col1, col2, col3, col4, col5 = st.columns([2.2, 1, 1, 1, 1.2])
                col1.write(f"**{user['username']}**")
                col2.write(f"ID: {user['id']}")
                col3.write(f"Role: **{user['role']}**")
                col4.write(f"Docs: {user['document_count']}")

                if user["id"] == current_user["id"]:
                    col5.caption("Current account")
                else:
                    new_role = "user" if user["role"] == "admin" else "admin"
                    label = "Make user" if new_role == "user" else "Make admin"
                    if col5.button(label, key=f"role_{user['id']}"):
                        try:
                            if st.session_state.db.set_user_role(user["id"], new_role):
                                st.success(f"{user['username']} is now {new_role}.")
                                st.rerun()
                            else:
                                st.warning("Role change was blocked. Keep at least one admin.")
                        except Exception as exc:
                            st.error(f"Role change failed: {exc}")

                st.divider()
        else:
            st.info("No users found.")

        st.subheader("➕ Create user")
        with st.form("admin_create_user"):
            admin_new_username = st.text_input("Username", key="admin_new_username")
            admin_new_password = st.text_input("Password", type="password", key="admin_new_password")
            admin_new_role = st.selectbox("Role", ["user", "admin"], key="admin_new_role")
            submitted = st.form_submit_button("Create account")
            if submitted:
                if not admin_new_username or not admin_new_password:
                    st.warning("Username and password are required.")
                else:
                    created = st.session_state.db.create_user(
                        admin_new_username, admin_new_password, admin_new_role
                    )
                    if created:
                        st.success(f"Created {created[1]} as {created[2]}.")
                        st.rerun()
                    else:
                        st.error("Could not create user. Username may already exist.")


# ============================================================
# SIDEBAR
# ============================================================

with st.sidebar:
    st.header("👤 Account")
    role_label = "🛡️ Admin" if is_admin else "👤 User"
    st.write(f"Logged in as **{current_user['username']}**")
    st.caption(f"Role: {role_label}")

    if is_admin:
        st.success("Admin controls available above.")

    if st.button("Logout", use_container_width=True):
        st.session_state.authenticated = False
        st.session_state.user = None
        st.session_state.messages = []
        st.session_state.conversation_id = None
        st.rerun()

    st.divider()

    # --------------------------------------------------------
    # CHUNKING: INGESTION SETTING
    # --------------------------------------------------------
    st.header("🧩 Chunking — ingestion")
    st.caption("Controls how a document is split before it is stored.")

    chunking_method = st.selectbox(
        "Chunking method",
        list(CHUNKING_METHODS.keys()),
        index=list(CHUNKING_METHODS.keys()).index("Multimodal"),
    )

    chunk_size = st.number_input(
        "Chunk size",
        min_value=100,
        max_value=20000,
        value=1000,
        step=100,
    )

    chunk_overlap = st.number_input(
        "Chunk overlap",
        min_value=0,
        max_value=5000,
        value=200,
        step=50,
    )

    if chunk_overlap >= chunk_size:
        st.error("Chunk overlap must be smaller than chunk size.")

    st.divider()

    # --------------------------------------------------------
    # RETRIEVAL: QUERY-TIME SETTING
    # --------------------------------------------------------
    st.header("🔎 Retrieval — query time")
    st.caption(f"pgvector: {'enabled' if PGVECTOR_ENABLED else 'disabled'} • embedding model: {EMBEDDING_MODEL}")
    st.caption("Controls how stored chunks are found when you ask a question.")

    retrieval_method = st.selectbox(
        "Retrieval method",
        [
            "PostgreSQL Full-Text",
            "PostgreSQL Full-Text + LIKE",
            "LIKE Keyword",
            "pgvector Semantic",
            "Hybrid (pgvector + PostgreSQL Full-Text)",
        ],
        index=1,
        help=(
            "Retrieval is independent of chunking. pgvector Semantic uses embedding similarity; "
            "Hybrid combines semantic similarity with PostgreSQL full-text ranking."
        ),
    )

    top_k = st.slider("Chunks to retrieve", 1, 20, 5)
    if retrieval_method in {"pgvector Semantic", "Hybrid (pgvector + PostgreSQL Full-Text)"}:
        st.caption(
            "pgvector is used for semantic similarity. If embeddings are not available, "
            "the application falls back to the existing PostgreSQL text retrieval instead of removing the feature."
        )

    st.divider()

    # --------------------------------------------------------
    # DOCUMENT UPLOAD
    # --------------------------------------------------------
    st.header("📁 Documents")
    uploaded_files = st.file_uploader(
        "Upload documents",
        type=["pdf", "txt", "md", "csv", "docx", "xlsx"],
        accept_multiple_files=True,
    )

    if uploaded_files and chunk_overlap < chunk_size:
        if st.button("⬆️ Process & Store Documents", use_container_width=True):
            progress = st.progress(0)
            total = len(uploaded_files)
            for index, uploaded_file in enumerate(uploaded_files, start=1):
                try:
                    extracted = extract_file(uploaded_file)
                    chunks, _ = create_chunks(
                        extracted, chunking_method, chunk_size, chunk_overlap
                    )
                    document_id = st.session_state.db.save_document(
                        user_id=current_user["id"],
                        filename=uploaded_file.name,
                        file_type=Path(uploaded_file.name).suffix.lower(),
                        file_bytes=uploaded_file.getvalue(),
                        chunking_method=chunking_method,
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap,
                        metadata=extracted.get("metadata", {}),
                    )
                    st.session_state.db.save_chunks(
                        document_id,
                        chunks,
                        embedder=generate_embedding if PGVECTOR_ENABLED else None,
                    )
                    progress.progress(index / total)
                    st.success(f"{uploaded_file.name}: {len(chunks)} chunks stored in PostgreSQL.")
                except Exception as exc:
                    st.session_state.db.rollback()
                    st.error(f"Failed to process {uploaded_file.name}: {exc}")

    st.divider()

    # --------------------------------------------------------
    # USER'S DOCUMENTS
    # --------------------------------------------------------
    try:
        documents = st.session_state.db.get_documents(current_user["id"], is_admin=False)
    except Exception as exc:
        st.session_state.db.rollback()
        documents = []
        st.error(f"Could not load documents: {exc}")

    st.session_state.documents = documents

    if documents:
        st.success(f"{len(documents)} document(s) in PostgreSQL.")
        document_options = {d["filename"]: d["id"] for d in documents}
        selected_document_names = st.multiselect(
            "Search only selected documents",
            list(document_options.keys()),
            default=list(document_options.keys()),
        )
        selected_document_ids = [document_options[n] for n in selected_document_names]
    else:
        selected_document_ids = []
        st.info("No documents uploaded yet.")

    st.divider()

    if st.button("🗑️ Clear conversation", use_container_width=True):
        st.session_state.messages = []
        try:
            st.session_state.conversation_id = st.session_state.db.create_conversation(
                current_user["id"], "New Conversation"
            )
        except Exception:
            st.session_state.db.rollback()
        st.rerun()


# ============================================================
# EMBEDDINGS
# ============================================================

def generate_embedding(text):
    """Return a fixed-size Gemini embedding when available.

    First uses a method exposed by the project's GeminiModel, then falls
    back to the official google-genai client. If embeddings are unavailable,
    callers can safely fall back to the existing PostgreSQL text retrieval.
    """
    if not text or not str(text).strip():
        return None

    llm = st.session_state.llm
    for name in ("embed_text", "get_embedding", "embedding", "embed"):
        method = getattr(llm, name, None)
        if callable(method):
            try:
                value = method(str(text))
                if hasattr(value, "values"):
                    value = value.values
                if isinstance(value, dict):
                    value = value.get("embedding", value.get("values"))
                if value is not None:
                    return [float(x) for x in value]
            except Exception:
                pass

    try:
        from google import genai
        api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
        if not api_key:
            return None
        client = genai.Client(api_key=api_key)
        result = client.models.embed_content(
            model=EMBEDDING_MODEL,
            contents=str(text),
            config={"output_dimensionality": PGVECTOR_DIMENSION},
        )
        emb = getattr(result, "embeddings", None)
        if emb:
            first = emb[0]
            values = getattr(first, "values", None)
            if values is not None:
                return [float(x) for x in values]
    except Exception:
        pass
    return None


def search_documents(query):
    return st.session_state.db.search_chunks(
        user_id=current_user["id"],
        query=query,
        selected_document_ids=selected_document_ids,
        chunk_types=None,
        top_k=top_k,
        retrieval_method=retrieval_method,
        query_embedding=(
            generate_embedding(query)
            if retrieval_method in {"pgvector Semantic", "Hybrid (pgvector + PostgreSQL Full-Text)"}
            else None
        ),
    )


def web_search(query):
    method = getattr(st.session_state.llm, "web_search", None)
    if not callable(method):
        return []
    try:
        results = method(query=query, max_results=top_k)
        if not isinstance(results, list):
            return []
        normalized = []
        for result in results:
            if not isinstance(result, dict):
                continue
            item = dict(result)
            item.setdefault("source_type", "web")
            item.setdefault("search_type", "web_search")
            item.setdefault("content", item.get("text", item.get("snippet", "")))
            normalized.append(item)
        return normalized
    except Exception as exc:
        st.warning(f"Web search failed: {exc}")
        return []


def generate_answer(query, chunks):
    generator = getattr(st.session_state.llm, "generate_answer", None)
    if not callable(generator):
        return chunks[0].get("content", "No answer available.") if chunks else "No answer available."
    try:
        return generator(
            query=query,
            chunks=chunks,
            source_type=(
                "web"
                if chunks and all(item.get("source_type") == "web" for item in chunks)
                else "document"
            ),
        )
    except Exception as exc:
        st.error(f"Answer generation failed: {exc}")
        return "I could not generate an answer from the retrieved evidence."


# ============================================================
# CHAT HISTORY
# ============================================================

for message in st.session_state.messages:
    role = message.get("role")
    content = message.get("content", "")
    if role in {"user", "assistant"}:
        with st.chat_message(role):
            st.markdown(content)


# ============================================================
# CHAT
# ============================================================

query = st.chat_input("Ask a question about your documents...")

if query:
    st.session_state.messages.append({"role": "user", "content": query})
    if st.session_state.conversation_id:
        try:
            st.session_state.db.save_message(st.session_state.conversation_id, "user", query)
        except Exception:
            st.session_state.db.rollback()

    with st.chat_message("user"):
        st.markdown(query)

    with st.chat_message("assistant"):
        try:
            with st.spinner(f"Retrieving using {retrieval_method}..."):
                document_results = search_documents(query)
        except Exception as exc:
            st.warning(f"Document retrieval failed: {exc}")
            st.session_state.db.rollback()
            document_results = []

        if document_results:
            results = document_results
            search_mode = "Documents"
        else:
            with st.spinner("No document evidence found. Searching web..."):
                results = web_search(query)
            search_mode = "Web"

        if results:
            with st.spinner(f"Generating answer from {search_mode}..."):
                answer = generate_answer(query, results)
        else:
            answer = (
                "I could not find sufficient evidence in the uploaded documents "
                "or available web search."
            )

        st.markdown(answer)

        col1, col2, col3 = st.columns(3)
        col1.metric("Retrieved results", len(results))
        col2.metric("Search", search_mode)
        col3.metric("Retrieval", retrieval_method if search_mode == "Documents" else "Web fallback")

        if results:
            with st.expander("📚 Retrieved Sources"):
                for index, source in enumerate(results, start=1):
                    source_name = (
                        source.get("filename")
                        or source.get("source")
                        or source.get("title")
                        or source.get("url")
                        or "Unknown"
                    )
                    chunk_type = source.get("chunk_type", "Web")
                    score = source.get("similarity_score", "")
                    page = source.get("page", "")

                    st.markdown(f"### {index}. {source_name}")
                    st.caption(f"Source type: {source.get('source_type', 'unknown')}")
                    st.caption(f"Chunk type: {chunk_type}")
                    if page:
                        st.caption(f"Page/Sheet: {page}")
                    if score != "":
                        st.caption(f"Relevance: {score}")

                    content = source.get("content") or source.get("text") or source.get("snippet") or ""
                    if content:
                        st.text(str(content)[:3000])

                    result_url = source.get("url")
                    if result_url:
                        try:
                            st.link_button("Open web source", result_url)
                        except Exception:
                            pass
                    st.divider()

    st.session_state.messages.append({"role": "assistant", "content": answer})
    if st.session_state.conversation_id:
        try:
            st.session_state.db.save_message(
                st.session_state.conversation_id, "assistant", answer
            )
        except Exception:
            st.session_state.db.rollback()
    st.rerun()


# ============================================================
# EMPTY STATE / EXPLANATION
# ============================================================

if not st.session_state.messages:
    st.markdown(
        """
        ## 👋 Welcome

        ### The important distinction

        **Chunking is an ingestion decision.** It decides how a document is split before storage.

        `Document → Extract → Chunk → PostgreSQL`

        **Retrieval is a query-time decision.** It decides how stored chunks are searched when the user asks a question.

        `Question → Retrieval → Relevant chunks → Gemini → Answer`

        These settings are intentionally separate.

        ### Current pipeline

        `Login → Upload → Extract → Chunk → Store in PostgreSQL → Retrieve → Web fallback → Gemini`

        ### PostgreSQL stores

        - Users and roles
        - Original document files
        - Documents and metadata
        - Document chunks
        - Chunking configuration
        - Conversations
        - Chat messages

        ### Roles

        - **Admin:** user/role management and system statistics.
        - **User:** their own documents and conversations.

        ### Retrieval available in this version

        - PostgreSQL Full-Text Search
        - PostgreSQL Full-Text + LIKE fallback
        - LIKE Keyword Search

        These are text retrieval methods. A true vector/semantic retrieval option should only be added when an embedding model and a vector index (for example pgvector) are actually configured.
        """
    )
